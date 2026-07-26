import io
import re
import openpyxl
from pypdf import PdfReader
from docx import Document as DocxDocument

from app.services.excel_engine.erp_excel_parser import ERPExcelParser

class DocumentParser:
    @staticmethod
    def extract_text_from_pdf(file_content: bytes) -> str:
        """Extract text from a PDF file using pypdf."""
        try:
            reader = PdfReader(io.BytesIO(file_content))
            text = []
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text.append(page_text)
            return "\n".join(text)
        except Exception as e:
            print(f"Error parsing PDF: {e}")
            raise ValueError(f"Failed to parse PDF document: {str(e)}")

    @staticmethod
    def parse_pothys_excel(file_content: bytes) -> dict:
        """
        Parse Pothys daily branch template using the new production format parser.
        """
        try:
            return ERPExcelParser.parse(file_content)
        except ValueError as ve:
            raise ve
        except Exception as e:
            raise ValueError(f"Failed to parse Excel document: {str(e)}")

    @staticmethod
    def extract_text_and_data_from_excel(file_content: bytes) -> tuple[dict, str]:
        """
        Extract structured metrics from the new production manager Excel template.
        """
        try:
            pothys_data = DocumentParser.parse_pothys_excel(file_content)
            sum_data = pothys_data["summary"]
            
            # sales_amount = total revenue aggregated from employee rows
            total_revenue = sum_data.get("total_revenue", 0.0)
            
            # Target achievement: compute from revenue if no explicit target
            target_achievement_val = 100.0
            if total_revenue:
                target_achievement_val = round((total_revenue / 500000.0) * 100, 2)

            # Populate metrics payload matching DailyReport columns
            extracted_metrics = {
                "sales_amount": total_revenue,
                "attendance_count": sum_data.get("employees_present", 0),
                "employees_present": sum_data.get("employees_present", 0),
                "employees_absent": sum_data.get("employees_absent", 0),
                "target_achievement": target_achievement_val,
                "remarks": sum_data.get("remarks", "None"),
                "issues": sum_data.get("operational_issues", "None"),
                "pothys_data": pothys_data  # pass full dict forward for DB insertion
            }
            
            full_text = (
                f"Pothys Swarna Mahal Daily Report\n"
                f"Date: {sum_data.get('report_date')}\n"
                f"Branch: {sum_data.get('branch_name')}\n"
                f"Gold: {sum_data.get('gold')}\n"
                f"Diamond: {sum_data.get('diamond')}\n"
                f"Platinum: {sum_data.get('platinum')}\n"
                f"Silver: {sum_data.get('silver')}\n"
                f"Silver MRP: {sum_data.get('silver_mrp')}\n"
                f"Total Revenue: {total_revenue}"
            )
            return extracted_metrics, full_text
        except ValueError as ve:
            if "Invalid report format" in str(ve):
                raise ve
            return {}, ""
        except Exception as e:
            print(f"Error parsing Excel: {e}")
            raise ValueError(f"Failed to parse Excel document: {str(e)}")

    @staticmethod
    def extract_text_from_docx(file_content: bytes) -> str:
        """Extract text from Word Document using python-docx."""
        try:
            doc = DocxDocument(io.BytesIO(file_content))
            text = []
            for paragraph in doc.paragraphs:
                if paragraph.text:
                    text.append(paragraph.text)
            
            # Extract table contents too
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text]
                    if row_text:
                        text.append(" | ".join(row_text))
            
            return "\n".join(text)
        except Exception as e:
            print(f"Error parsing Word Document: {e}")
            raise ValueError(f"Failed to parse Word Document: {str(e)}")

    @classmethod
    def parse_document(cls, file_content: bytes, file_name: str) -> tuple[dict, str]:
        """
        Main routing method to parse a document based on its extension.
        Returns a tuple (extracted_metrics, full_text_content).
        """
        ext = file_name.split(".")[-1].lower()
        full_text = ""
        extracted_metrics = {}

        if ext == "pdf":
            full_text = cls.extract_text_from_pdf(file_content)
        elif ext in ["xlsx", "xls"]:
            extracted_metrics, full_text = cls.extract_text_and_data_from_excel(file_content)
        elif ext in ["docx", "doc"]:
            full_text = cls.extract_text_from_docx(file_content)
        else:
            raise ValueError(f"Unsupported file format: {ext}")

        # Run heuristic regex parses on PDF/Word text if structured data is not yet found
        if "sales_amount" not in extracted_metrics:
            sales_match = re.search(r"(?:sales|revenue|collection)[:\s]*Rs\.?\s*([\d,]+(?:\.\d{2})?)", full_text, re.IGNORECASE)
            if sales_match:
                extracted_metrics["sales_amount"] = float(sales_match.group(1).replace(",", ""))
        
        if "attendance_count" not in extracted_metrics:
            att_match = re.search(r"(?:attendance|staff present|headcount)[:\s]*(\d+)", full_text, re.IGNORECASE)
            if att_match:
                extracted_metrics["attendance_count"] = int(att_match.group(1))

        if "target_achievement" not in extracted_metrics:
            target_match = re.search(r"(?:target achievement|achievement|target)[:\s]*(\d+(?:\.\d+)?)\s*%", full_text, re.IGNORECASE)
            if target_match:
                extracted_metrics["target_achievement"] = float(target_match.group(1))

        # Scan for remarks / issues
        remarks_match = re.search(r"(?:remarks|notes|comments)[:\s]*(.*)", full_text, re.IGNORECASE)
        if remarks_match:
            extracted_metrics["remarks"] = remarks_match.group(1).strip()[:1000]

        issues_match = re.search(r"(?:issues|complaints|problems|incidents)[:\s]*(.*)", full_text, re.IGNORECASE)
        if issues_match:
            extracted_metrics["issues"] = issues_match.group(1).strip()[:1000]

        return extracted_metrics, full_text

document_parser = DocumentParser()
